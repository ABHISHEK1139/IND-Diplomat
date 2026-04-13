"""
Report Generator for IND-Diplomat
Generates PDF and DOCX briefing documents from analysis results.
"""

import os
from typing import Dict, Any, List, Optional
from datetime import datetime
from io import BytesIO

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    print("[ReportGen] Warning: reportlab not installed. PDF generation disabled.")

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[ReportGen] Warning: python-docx not installed. DOCX generation disabled.")


class ReportGenerator:
    """
    Production-grade report generator with:
    1. PDF generation (ReportLab)
    2. DOCX generation (python-docx)
    3. Briefing document templates
    4. Custom branding support
    """
    
    def __init__(self):
        self.output_dir = os.getenv("REPORT_OUTPUT_DIR", "./reports")
        os.makedirs(self.output_dir, exist_ok=True)
        
        self.org_name = "IND-Diplomat Intelligence System"
        self.classification = "CONFIDENTIAL"
    
    def generate_pdf(self, report_data: Dict[str, Any]) -> Optional[bytes]:
        """Generates PDF briefing document."""
        if not REPORTLAB_AVAILABLE:
            print("[ReportGen] PDF generation not available")
            return None
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.75*inch, bottomMargin=0.75*inch)
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#1a237e'),
            spaceAfter=12
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#303f9f'),
            spaceBefore=12,
            spaceAfter=6
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            leading=14,
            spaceAfter=10
        )
        
        elements = []
        
        # Header
        elements.append(Paragraph(self.org_name, title_style))
        elements.append(Paragraph(f"Classification: {self.classification}", styles['Normal']))
        elements.append(Paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}", styles['Normal']))
        elements.append(Spacer(1, 0.25*inch))
        
        # Query Section
        if report_data.get("query"):
            elements.append(Paragraph("Query", heading_style))
            elements.append(Paragraph(report_data["query"], body_style))
        
        # Analysis Section
        if report_data.get("answer"):
            elements.append(Paragraph("Analysis", heading_style))
            # Split into paragraphs
            answer_paragraphs = report_data["answer"].split('\n\n')
            for para in answer_paragraphs:
                if para.strip():
                    elements.append(Paragraph(para.strip(), body_style))
        
        # Sources Section
        if report_data.get("sources"):
            elements.append(Paragraph("Sources", heading_style))
            for i, source in enumerate(report_data["sources"][:5], 1):
                source_text = source if isinstance(source, str) else str(source)
                elements.append(Paragraph(f"{i}. {source_text[:200]}...", body_style))
        
        # Verification Section
        if report_data.get("faithfulness_score") is not None:
            elements.append(Paragraph("Verification", heading_style))
            score = report_data["faithfulness_score"]
            elements.append(Paragraph(f"Faithfulness Score: {score:.2%}", body_style))
            
            if report_data.get("warnings"):
                elements.append(Paragraph(f"Warnings: {', '.join(report_data['warnings'])}", body_style))
        
        # Provenance Section
        if report_data.get("c2pa_manifest"):
            elements.append(Paragraph("Digital Provenance", heading_style))
            manifest = report_data["c2pa_manifest"]
            if manifest.get("signature"):
                elements.append(Paragraph(f"Signed with: {manifest['signature'].get('algorithm', 'N/A')}", body_style))
        
        # Build PDF
        doc.build(elements)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        
        return pdf_bytes
    
    def generate_docx(self, report_data: Dict[str, Any]) -> Optional[bytes]:
        """Generates DOCX briefing document."""
        if not DOCX_AVAILABLE:
            print("[ReportGen] DOCX generation not available")
            return None
        
        doc = Document()
        
        # Title
        title = doc.add_heading(self.org_name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        meta = doc.add_paragraph()
        meta.add_run(f"Classification: {self.classification}\n").bold = True
        meta.add_run(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
        
        doc.add_paragraph()
        
        # Query
        if report_data.get("query"):
            doc.add_heading("Query", level=1)
            doc.add_paragraph(report_data["query"])
        
        # Analysis
        if report_data.get("answer"):
            doc.add_heading("Analysis", level=1)
            doc.add_paragraph(report_data["answer"])
        
        # Sources
        if report_data.get("sources"):
            doc.add_heading("Sources", level=1)
            for i, source in enumerate(report_data["sources"][:5], 1):
                source_text = source if isinstance(source, str) else str(source)
                doc.add_paragraph(f"{i}. {source_text[:200]}...", style='List Number')
        
        # Verification
        if report_data.get("faithfulness_score") is not None:
            doc.add_heading("Verification", level=1)
            doc.add_paragraph(f"Faithfulness Score: {report_data['faithfulness_score']:.2%}")
        
        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()
        
        return docx_bytes
    
    def save_report(self, report_data: Dict[str, Any], filename: str, format: str = "pdf") -> Optional[str]:
        """Saves report to file."""
        if format == "pdf":
            content = self.generate_pdf(report_data)
            ext = ".pdf"
        elif format == "docx":
            content = self.generate_docx(report_data)
            ext = ".docx"
        else:
            print(f"[ReportGen] Unsupported format: {format}")
            return None
        
        if content is None:
            return None
        
        filepath = os.path.join(self.output_dir, f"{filename}{ext}")
        with open(filepath, "wb") as f:
            f.write(content)
        
        print(f"[ReportGen] Saved report to: {filepath}")
        return filepath


# Singleton instance
report_generator = ReportGenerator()
